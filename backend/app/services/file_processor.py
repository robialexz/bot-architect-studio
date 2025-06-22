"""
File processing service for FlowsyAI Backend
Handles various file types and extracts content/metadata
"""

import os
import json
import csv
import io
from pathlib import Path
from typing import Dict, Any, Optional, List
from datetime import datetime

import pandas as pd
import PyPDF2
import docx
from PIL import Image, ExifTags
try:
    import eyed3
except ImportError:
    eyed3 = None

try:
    import magic
except ImportError:
    magic = None

import chardet

from app.core.logging import get_logger

logger = get_logger(__name__)

class FileProcessor:
    """Service for processing different file types"""
    
    def __init__(self):
        self.processors = {
            'text': self._process_text_file,
            'document': self._process_document,
            'image': self._process_image,
            'audio': self._process_audio,
            'video': self._process_video,
            'data': self._process_data_file,
            'code': self._process_code_file
        }
    
    async def process_file(self, file_path: str, category: str, options: Dict[str, Any] = None) -> Dict[str, Any]:
        """Process file based on its category"""
        try:
            if category not in self.processors:
                return {
                    'success': False,
                    'error': f'Unsupported file category: {category}'
                }
            
            processor = self.processors[category]
            result = await processor(file_path, options or {})
            
            # Add common metadata
            result['processed_at'] = datetime.utcnow().isoformat()
            result['file_path'] = file_path
            result['category'] = category
            
            return result
            
        except Exception as e:
            logger.error(f"File processing failed for {file_path}: {e}")
            return {
                'success': False,
                'error': str(e),
                'processed_at': datetime.utcnow().isoformat()
            }
    
    async def _process_text_file(self, file_path: str, options: Dict[str, Any]) -> Dict[str, Any]:
        """Process text files (txt, md, csv, json, etc.)"""
        try:
            file_ext = Path(file_path).suffix.lower()
            
            # Detect encoding
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                encoding_result = chardet.detect(raw_data)
                encoding = encoding_result.get('encoding', 'utf-8')
            
            # Read file content
            with open(file_path, 'r', encoding=encoding) as f:
                content = f.read()
            
            result = {
                'success': True,
                'content': content,
                'encoding': encoding,
                'character_count': len(content),
                'line_count': len(content.splitlines()),
                'word_count': len(content.split()),
                'file_type': file_ext
            }
            
            # Special processing for specific file types
            if file_ext == '.csv':
                result.update(await self._analyze_csv(file_path))
            elif file_ext == '.json':
                result.update(await self._analyze_json(content))
            elif file_ext in ['.md', '.markdown']:
                result.update(await self._analyze_markdown(content))
            
            return result
            
        except Exception as e:
            return {
                'success': False,
                'error': f'Text processing failed: {str(e)}'
            }
    
    async def _process_document(self, file_path: str, options: Dict[str, Any]) -> Dict[str, Any]:
        """Process document files (PDF, DOCX, etc.)"""
        try:
            file_ext = Path(file_path).suffix.lower()
            
            if file_ext == '.pdf':
                return await self._process_pdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                return await self._process_docx(file_path)
            else:
                return {
                    'success': False,
                    'error': f'Unsupported document format: {file_ext}'
                }
                
        except Exception as e:
            return {
                'success': False,
                'error': f'Document processing failed: {str(e)}'
            }
    
    async def _process_image(self, file_path: str, options: Dict[str, Any]) -> Dict[str, Any]:
        """Process image files"""
        try:
            with Image.open(file_path) as img:
                # Basic image info
                result = {
                    'success': True,
                    'width': img.width,
                    'height': img.height,
                    'format': img.format,
                    'mode': img.mode,
                    'size_bytes': os.path.getsize(file_path)
                }
                
                # Extract EXIF data
                if hasattr(img, '_getexif'):
                    exif_data = img._getexif()
                    if exif_data:
                        exif = {}
                        for tag_id, value in exif_data.items():
                            tag = ExifTags.TAGS.get(tag_id, tag_id)
                            exif[tag] = str(value)
                        result['exif'] = exif
                
                # Color analysis
                if img.mode == 'RGB':
                    colors = img.getcolors(maxcolors=256*256*256)
                    if colors:
                        dominant_color = max(colors, key=lambda x: x[0])[1]
                        result['dominant_color'] = dominant_color
                
                # Generate thumbnail if requested
                if options.get('generate_thumbnail', False):
                    thumbnail_path = await self._generate_thumbnail(file_path)
                    result['thumbnail_path'] = thumbnail_path
                
                return result
                
        except Exception as e:
            return {
                'success': False,
                'error': f'Image processing failed: {str(e)}'
            }
    
    async def _process_audio(self, file_path: str, options: Dict[str, Any]) -> Dict[str, Any]:
        """Process audio files"""
        try:
            audiofile = eyed3.load(file_path)
            
            result = {
                'success': True,
                'size_bytes': os.path.getsize(file_path)
            }
            
            if audiofile and audiofile.info:
                result.update({
                    'duration_seconds': audiofile.info.time_secs,
                    'bitrate': audiofile.info.bit_rate[1] if audiofile.info.bit_rate else None,
                    'sample_rate': audiofile.info.sample_freq,
                    'channels': audiofile.info.mode
                })
            
            if audiofile and audiofile.tag:
                result.update({
                    'title': audiofile.tag.title,
                    'artist': audiofile.tag.artist,
                    'album': audiofile.tag.album,
                    'genre': str(audiofile.tag.genre) if audiofile.tag.genre else None,
                    'year': audiofile.tag.getBestDate()
                })
            
            return result
            
        except Exception as e:
            return {
                'success': False,
                'error': f'Audio processing failed: {str(e)}'
            }
    
    async def _process_video(self, file_path: str, options: Dict[str, Any]) -> Dict[str, Any]:
        """Process video files"""
        try:
            # Basic file info
            result = {
                'success': True,
                'size_bytes': os.path.getsize(file_path)
            }
            
            # For more detailed video processing, you would use libraries like ffmpeg-python
            # This is a basic implementation
            
            return result
            
        except Exception as e:
            return {
                'success': False,
                'error': f'Video processing failed: {str(e)}'
            }
    
    async def _process_data_file(self, file_path: str, options: Dict[str, Any]) -> Dict[str, Any]:
        """Process data files (Excel, CSV, etc.)"""
        try:
            file_ext = Path(file_path).suffix.lower()
            
            if file_ext in ['.xlsx', '.xls']:
                return await self._process_excel(file_path)
            elif file_ext == '.csv':
                return await self._analyze_csv(file_path)
            else:
                return {
                    'success': False,
                    'error': f'Unsupported data format: {file_ext}'
                }
                
        except Exception as e:
            return {
                'success': False,
                'error': f'Data processing failed: {str(e)}'
            }
    
    async def _process_code_file(self, file_path: str, options: Dict[str, Any]) -> Dict[str, Any]:
        """Process code files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            file_ext = Path(file_path).suffix.lower()
            
            result = {
                'success': True,
                'content': content,
                'line_count': len(content.splitlines()),
                'character_count': len(content),
                'language': self._detect_language(file_ext),
                'file_type': file_ext
            }
            
            # Basic code analysis
            lines = content.splitlines()
            result.update({
                'blank_lines': len([line for line in lines if not line.strip()]),
                'comment_lines': self._count_comment_lines(lines, file_ext),
                'code_lines': len([line for line in lines if line.strip() and not self._is_comment_line(line, file_ext)])
            })
            
            return result
            
        except Exception as e:
            return {
                'success': False,
                'error': f'Code processing failed: {str(e)}'
            }
    
    # Helper methods
    
    async def _process_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract text from PDF"""
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text()
                
                return {
                    'success': True,
                    'content': text,
                    'page_count': len(pdf_reader.pages),
                    'character_count': len(text),
                    'word_count': len(text.split())
                }
                
        except Exception as e:
            return {
                'success': False,
                'error': f'PDF processing failed: {str(e)}'
            }
    
    async def _process_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract text from DOCX"""
        try:
            doc = docx.Document(file_path)
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return {
                'success': True,
                'content': text,
                'paragraph_count': len(doc.paragraphs),
                'character_count': len(text),
                'word_count': len(text.split())
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f'DOCX processing failed: {str(e)}'
            }
    
    async def _analyze_csv(self, file_path: str) -> Dict[str, Any]:
        """Analyze CSV file structure"""
        try:
            df = pd.read_csv(file_path)
            
            return {
                'csv_analysis': {
                    'row_count': len(df),
                    'column_count': len(df.columns),
                    'columns': df.columns.tolist(),
                    'data_types': df.dtypes.astype(str).to_dict(),
                    'null_counts': df.isnull().sum().to_dict(),
                    'sample_data': df.head().to_dict('records')
                }
            }
            
        except Exception as e:
            return {
                'csv_analysis': {
                    'error': f'CSV analysis failed: {str(e)}'
                }
            }
    
    async def _analyze_json(self, content: str) -> Dict[str, Any]:
        """Analyze JSON structure"""
        try:
            data = json.loads(content)
            
            def analyze_structure(obj, depth=0):
                if depth > 5:  # Prevent infinite recursion
                    return "max_depth_reached"
                
                if isinstance(obj, dict):
                    return {key: analyze_structure(value, depth + 1) for key, value in obj.items()}
                elif isinstance(obj, list):
                    if obj:
                        return [analyze_structure(obj[0], depth + 1)]
                    return []
                else:
                    return type(obj).__name__
            
            return {
                'json_analysis': {
                    'structure': analyze_structure(data),
                    'is_array': isinstance(data, list),
                    'is_object': isinstance(data, dict),
                    'size': len(str(data))
                }
            }
            
        except Exception as e:
            return {
                'json_analysis': {
                    'error': f'JSON analysis failed: {str(e)}'
                }
            }
    
    async def _analyze_markdown(self, content: str) -> Dict[str, Any]:
        """Analyze Markdown structure"""
        lines = content.splitlines()
        
        headers = []
        for line in lines:
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                text = line.lstrip('#').strip()
                headers.append({'level': level, 'text': text})
        
        return {
            'markdown_analysis': {
                'header_count': len(headers),
                'headers': headers,
                'has_code_blocks': '```' in content,
                'has_links': '[' in content and '](' in content,
                'has_images': '![' in content
            }
        }
    
    async def _process_excel(self, file_path: str) -> Dict[str, Any]:
        """Process Excel file"""
        try:
            xl_file = pd.ExcelFile(file_path)
            
            sheets_info = {}
            for sheet_name in xl_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                sheets_info[sheet_name] = {
                    'row_count': len(df),
                    'column_count': len(df.columns),
                    'columns': df.columns.tolist()
                }
            
            return {
                'success': True,
                'excel_analysis': {
                    'sheet_count': len(xl_file.sheet_names),
                    'sheet_names': xl_file.sheet_names,
                    'sheets_info': sheets_info
                }
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f'Excel processing failed: {str(e)}'
            }
    
    async def _generate_thumbnail(self, file_path: str) -> str:
        """Generate thumbnail for image"""
        try:
            with Image.open(file_path) as img:
                img.thumbnail((200, 200))
                
                thumbnail_path = file_path.replace(Path(file_path).suffix, '_thumb.jpg')
                img.save(thumbnail_path, 'JPEG')
                
                return thumbnail_path
                
        except Exception as e:
            logger.error(f"Thumbnail generation failed: {e}")
            return None
    
    def _detect_language(self, file_ext: str) -> str:
        """Detect programming language from file extension"""
        language_map = {
            '.py': 'python',
            '.js': 'javascript',
            '.ts': 'typescript',
            '.html': 'html',
            '.css': 'css',
            '.sql': 'sql',
            '.r': 'r',
            '.ipynb': 'jupyter',
            '.java': 'java',
            '.cpp': 'cpp',
            '.c': 'c',
            '.php': 'php',
            '.rb': 'ruby',
            '.go': 'go',
            '.rs': 'rust'
        }
        return language_map.get(file_ext, 'unknown')
    
    def _count_comment_lines(self, lines: List[str], file_ext: str) -> int:
        """Count comment lines in code"""
        comment_count = 0
        for line in lines:
            if self._is_comment_line(line.strip(), file_ext):
                comment_count += 1
        return comment_count
    
    def _is_comment_line(self, line: str, file_ext: str) -> bool:
        """Check if line is a comment"""
        if not line:
            return False
        
        if file_ext in ['.py', '.r']:
            return line.startswith('#')
        elif file_ext in ['.js', '.ts', '.java', '.cpp', '.c', '.php', '.go', '.rs']:
            return line.startswith('//') or line.startswith('/*')
        elif file_ext in ['.html', '.xml']:
            return line.startswith('<!--')
        elif file_ext == '.css':
            return line.startswith('/*')
        elif file_ext == '.sql':
            return line.startswith('--')
        
        return False
